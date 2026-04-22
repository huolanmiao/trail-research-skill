"""Shared utilities for Word document generation."""

import platform
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _get_fonts():
    """Return (body_font, heading_font) suitable for the current OS."""
    system = platform.system()
    if system == "Darwin":
        return ("Songti SC", "Heiti SC")
    elif system == "Windows":
        return ("宋体", "黑体")
    else:
        return ("Noto Serif CJK SC", "Noto Sans CJK SC")


_body_font, _heading_font = _get_fonts()


def create_document() -> Document:
    """Create a new Document with default Chinese-friendly styling."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = _body_font
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), _body_font)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    return doc


def add_heading_styled(doc, text, level=1):
    """Add a heading with Chinese-friendly font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = _heading_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), _heading_font)
    return heading


def add_paragraph_styled(doc, text, bold=False, font_size=None, alignment=None, font_name=None):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    fn = font_name or _body_font
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        para.alignment = alignment
    return para


def set_cell_text(cell, text, bold=False, font_size=10.5, alignment=None, font_name=None):
    """Set cell content with formatting. Clears all existing paragraphs."""
    fn = font_name or _body_font
    # Remove all paragraphs after the first, then clear the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(str(text))
    run.font.name = fn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        para.alignment = alignment


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_with_data(doc, headers, rows, col_widths=None, header_color="D9E2F3"):
    """Create a formatted table with borders and header shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, font_size=10.5)
        set_cell_shading(cell, header_color)

    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            set_cell_text(cell, val)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_field_table(doc, fields, col_widths=None):
    """Add a 2-column key-value table (field name | field value)."""
    headers = ["项目", "内容"]
    rows = [[k, v] for k, v in fields]
    return add_table_with_data(doc, headers, rows, col_widths or [4, 12])


def add_signature_block(doc, lines):
    """Add approval/signature section at document end."""
    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        run.font.name = _body_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), _body_font)
        run.font.size = Pt(10.5)


def add_bullet_list(doc, items, indent_level=0):
    """Add a bulleted list."""
    for item in items:
        para = doc.add_paragraph(style='List Bullet')
        para.clear()
        run = para.add_run(item)
        run.font.name = _body_font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), _body_font)
        run.font.size = Pt(10.5)


def save_document(doc, filepath):
    """Save document to filepath. Raises OSError on failure."""
    doc.save(filepath)
    return filepath
